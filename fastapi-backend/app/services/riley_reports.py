import asyncio
import io
import json
import logging
import re
import uuid
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

from fastapi.concurrency import run_in_threadpool
from google.api_core.exceptions import AlreadyExists
from google.cloud import tasks_v2
from qdrant_client.http.models import PointStruct
from qdrant_client.http.models import FieldCondition, Filter, MatchValue

from app.core.config import get_settings
from app.core.personas import get_persona_context
from app.services.genai_client import get_genai_client
from app.services.graph import GraphService
from app.services.provider_fallback import (
    classify_gemini_generation_failure,
    classify_openai_generation_failure,
    generate_text_with_gemini,
)
from app.services.qdrant import vector_service
from app.services.rerank import rerank_candidates
from app.services.storage import StorageService
from app.services.token_utils import estimate_tokens

logger = logging.getLogger(__name__)


def _normalize_report_mode(mode: Optional[str]) -> str:
    if (mode or "").strip().lower() == "normal":
        return "normal"
    return "deep"


def _normalize_report_type(report_type: Optional[str]) -> str:
    normalized = (report_type or "").strip().lower()
    allowed = {
        "summary",
        "strategy_memo",
        "audience_analysis",
        "narrative_brief",
        "opposition_framing_brief",
    }
    if normalized in allowed:
        return normalized
    return "strategy_memo"


def _report_type_instruction(report_type: str) -> str:
    normalized = _normalize_report_type(report_type)
    if normalized == "summary":
        return """REPORT TYPE: summary
Purpose:
- Produce a neutral analytical synthesis of campaign documents without memo framing.
Prioritize:
- Clear overview of the campaign narrative across materials.
- Key recurring themes and evidence-backed patterns.
- Contradictions, tensions, and points of narrative drift across sources.
- Potential vulnerabilities that opponents could exploit.
- Concrete evidence references for major claims.
Constraints:
- Do not use memo headers (no TO / FROM / SUBJECT).
- Do not include strategic recommendations unless the user explicitly asks for them.
Output structure:
- Narrative Overview
- Key Themes Across Documents
- Contradictions and Tensions
- Potential Opponent-Exploitable Vulnerabilities
- Evidence References"""
    if normalized == "audience_analysis":
        return """REPORT TYPE: audience_analysis
Purpose:
- Explain how key audience segments/personas are likely to interpret campaign messaging and where persuasion potential is strongest or weakest.
Prioritize:
- Segment-by-segment motivations, barriers, trust drivers, and likely emotional/cognitive triggers.
- Evidence-backed resonance differences across audiences.
- Practical implications for persuasion sequencing, targeting, and message tailoring.
Strategic questions to answer:
- Which audiences are most movable right now, and why?
- What message themes are likely to land or backfire by segment?
- What audience-specific adjustments should the campaign make immediately?
Output structure:
- Audience Segments and Persona Snapshot
- Motivations and Barriers by Segment
- Resonance and Friction by Message Theme
- Persuasion Implications and Priority Segments
- Audience-Specific Recommendations and Tests"""
    if normalized == "narrative_brief":
        return """REPORT TYPE: narrative_brief
Purpose:
- Identify dominant narratives in the corpus and recommend a disciplined narrative architecture for campaign execution.
Prioritize:
- Narrative patterns repeated across sources.
- Framing opportunities that can unify campaign communication.
- Message discipline rules that reduce drift and contradiction.
Strategic questions to answer:
- What narratives currently dominate and which are emerging?
- Which narratives should the campaign amplify, refine, or retire?
- What disciplined message architecture should guide spokespeople and content?
Output structure:
- Dominant Narratives in the Corpus
- Narrative Opportunities and Gaps
- Recommended Core Frame and Supporting Frames
- Message Discipline Guardrails
- Immediate Narrative Rollout Plan"""
    if normalized == "opposition_framing_brief":
        return """REPORT TYPE: opposition_framing_brief
Purpose:
- Map opposition attack narratives and vulnerabilities, then provide concrete counter-framing strategy.
Prioritize:
- Documented attack lines, hostile framings, and likely amplification vectors.
- Campaign vulnerabilities most exposed to opposition pressure.
- Counter-frame options with tradeoffs and execution guidance.
Strategic questions to answer:
- Which opposition narratives are most dangerous and why?
- Where is the campaign most vulnerable to framing attacks?
- What counter-framing moves are highest-value in the near term?
Output structure:
- Opposition Narrative Map
- Vulnerability Assessment
- Counter-Framing Options and Tradeoffs
- Defensive and Offensive Messaging Recommendations
- Rapid Response and Monitoring Plan"""
    return """REPORT TYPE: strategy_memo
Purpose:
- Produce a high-confidence strategic memo for campaign decision-making.
Prioritize:
- Campaign objective clarity, narrative opportunities, and messaging recommendations.
- Risks, tradeoffs, and execution implications.
- Immediate next steps that can be operationalized by the team.
Strategic questions to answer:
- What objective should guide near-term strategy and why?
- Which narrative and message opportunities are strongest in current evidence?
- What are the key risks/tradeoffs and what should we do next?
Output structure:
- Campaign Objective and Strategic Context
- Narrative and Messaging Opportunities
- Recommended Message Strategy
- Risks and Tradeoffs
- Next Steps and Execution Priorities"""


def _derive_report_title(query: str, explicit_title: Optional[str]) -> str:
    if explicit_title and explicit_title.strip():
        return explicit_title.strip()[:180]
    normalized = (query or "").strip()
    if not normalized:
        return "Riley Strategy Report"
    if len(normalized) <= 120:
        return normalized
    return f"{normalized[:117].rstrip()}..."


def _candidate_id(result: Dict[str, Any]) -> str:
    payload = result.get("payload", {}) or {}
    chunk_id = payload.get("chunk_id")
    if chunk_id:
        return str(chunk_id)
    parent = payload.get("parent_file_id")
    idx = payload.get("chunk_index")
    if parent is not None and idx is not None:
        return f"{parent}::chunk::{idx}"
    return str(result.get("id", ""))


def _apply_rerank_order(
    private_results: List[Dict[str, Any]],
    global_results: List[Dict[str, Any]],
    ranked_ids: List[str],
) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    if not ranked_ids:
        return private_results, global_results

    by_id: Dict[str, Tuple[str, Dict[str, Any]]] = {}
    for item in private_results:
        by_id[_candidate_id(item)] = ("private", item)
    for item in global_results:
        by_id[_candidate_id(item)] = ("global", item)

    ordered_private: List[Dict[str, Any]] = []
    ordered_global: List[Dict[str, Any]] = []
    used: set[str] = set()
    for candidate_id in ranked_ids:
        if candidate_id in used:
            continue
        entry = by_id.get(candidate_id)
        if not entry:
            continue
        used.add(candidate_id)
        scope, item = entry
        if scope == "private":
            ordered_private.append(item)
        else:
            ordered_global.append(item)

    for item in private_results:
        cid = _candidate_id(item)
        if cid not in used:
            ordered_private.append(item)
    for item in global_results:
        cid = _candidate_id(item)
        if cid not in used:
            ordered_global.append(item)
    return ordered_private, ordered_global


def _get_text_for_rag(payload: Dict[str, Any], settings: Any) -> str:
    file_type = payload.get("file_type") or payload.get("type", "")
    is_image = file_type.lower() in ("png", "jpg", "jpeg", "webp", "tiff")
    if is_image:
        ai_enabled = payload.get("ai_enabled", False)
        ocr_status = payload.get("ocr_status", "not_requested")
        ocr_confidence = payload.get("ocr_confidence")
        ocr_text = payload.get("ocr_text")
        if (
            ai_enabled
            and ocr_status == "complete"
            and ocr_text
            and (ocr_confidence is None or ocr_confidence >= settings.OCR_MIN_CONFIDENCE)
        ):
            return ocr_text
    return payload.get("content") or payload.get("content_preview") or ""


def _format_rag_context(
    private_results: List[Dict[str, Any]],
    global_results: List[Dict[str, Any]],
    graph_results: str,
    file_manifest: List[str],
    *,
    max_text_chars_per_result: int = 2200,
) -> str:
    settings = get_settings()
    parts: List[str] = []
    if graph_results:
        parts.append("=== KNOWLEDGE GRAPH ===\n" + graph_results)

    if file_manifest:
        manifest_text = "\n".join(f"- {name}" for name in file_manifest[:80])
        parts.append("=== FILE MANIFEST ===\n" + manifest_text)

    if private_results:
        private_lines: List[str] = []
        for idx, result in enumerate(private_results, start=1):
            payload = result.get("payload", {}) or {}
            filename = payload.get("filename", "Unknown")
            location = payload.get("location_value") or payload.get("location") or "Unknown location"
            text = _get_text_for_rag(payload, settings)
            if len(text) > max_text_chars_per_result:
                text = f"{text[:max_text_chars_per_result].rstrip()}…"
            private_lines.append(
                f"[Private {idx}] {filename} | {location}\n{text}"
            )
        parts.append("=== CAMPAIGN DOCUMENT ARCHIVE ===\n" + "\n\n".join(private_lines))

    if global_results:
        global_lines: List[str] = []
        for idx, result in enumerate(global_results, start=1):
            payload = result.get("payload", {}) or {}
            filename = payload.get("filename", "Unknown")
            location = payload.get("location_value") or payload.get("location") or "Unknown location"
            text = _get_text_for_rag(payload, settings)
            if len(text) > max_text_chars_per_result:
                text = f"{text[:max_text_chars_per_result].rstrip()}…"
            global_lines.append(
                f"[Global {idx}] {filename} | {location}\n{text}"
            )
        parts.append("=== GLOBAL KNOWLEDGE ARCHIVE ===\n" + "\n\n".join(global_lines))

    return "\n\n".join(parts)


def _safe_json_loads(raw: Any, default: Any) -> Any:
    if not isinstance(raw, str) or not raw.strip():
        return default
    try:
        return json.loads(raw)
    except Exception:
        return default


async def _load_parent_document_intelligence(
    *,
    collection_name: str,
    results: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    parent_ids: List[str] = []
    seen: set[str] = set()
    for result in results:
        payload = result.get("payload", {}) or {}
        parent_id = str(payload.get("parent_file_id") or "").strip()
        if not parent_id or parent_id in seen:
            continue
        seen.add(parent_id)
        parent_ids.append(parent_id)
    if not parent_ids:
        return []
    try:
        parent_points = await vector_service.client.retrieve(
            collection_name=collection_name,
            ids=parent_ids,
            with_payload=True,
            with_vectors=False,
        )
    except Exception:
        return []
    docs: List[Dict[str, Any]] = []
    for point in parent_points:
        payload = point.payload or {}
        if str(payload.get("analysis_status") or "").lower() != "complete":
            continue
        docs.append(payload)
    return docs


def _build_doc_intel_context_block(doc_intel_items: List[Dict[str, Any]]) -> str:
    if not doc_intel_items:
        return ""
    lines: List[str] = ["=== SYNTHESIZED DOCUMENT INTELLIGENCE ==="]
    for item in doc_intel_items[:14]:
        filename = str(item.get("filename") or "Unknown")
        short_summary = str(item.get("doc_summary_short") or "").strip()
        themes = item.get("key_themes") or []
        tones = item.get("tone_labels") or []
        framings = item.get("framing_labels") or []
        opportunities = item.get("strategic_opportunities") or []
        risks = item.get("persuasion_risks") or []
        fidelity = str(item.get("analysis_fidelity_level") or "unknown").strip()
        exec_mode = str(item.get("analysis_execution_mode") or "unknown").strip()
        chunks_cov = item.get("analysis_chunks_coverage_ratio")
        chars_cov = item.get("analysis_chars_coverage_ratio")
        bands_total = int(item.get("analysis_total_bands") or 0)
        bands_analyzed = int(item.get("analysis_analyzed_bands") or 0)
        band_cov = item.get("analysis_band_coverage_ratio")
        validation_status = str(item.get("analysis_validation_status") or "").strip()
        validation_note = str(item.get("analysis_validation_note") or "").strip()
        contradiction_count = int(item.get("analysis_contradiction_count") or 0)
        failed_bands = int(item.get("analysis_failed_bands_count") or 0)
        high_signal_cov = item.get("analysis_high_signal_band_coverage_ratio")
        appendix_required = bool(item.get("analysis_appendix_required"))
        appendix_covered = bool(item.get("analysis_appendix_covered"))
        lines.append(f"[Doc Intelligence] {filename}")
        lines.append(f"- fidelity: {fidelity}")
        lines.append(f"- execution_mode: {exec_mode}")
        if bands_total > 0:
            if band_cov is not None:
                lines.append(
                    f"- band_coverage: {bands_analyzed}/{bands_total} ({float(band_cov or 0.0):.2%})"
                )
            else:
                lines.append(f"- band_coverage: {bands_analyzed}/{bands_total}")
        if chunks_cov is not None or chars_cov is not None:
            lines.append(
                f"- coverage: chunks={float(chunks_cov or 0.0):.2%}, chars={float(chars_cov or 0.0):.2%}"
            )
        if validation_status:
            lines.append(f"- validation_status: {validation_status}")
        if validation_note:
            lines.append(f"- validation_note: {validation_note}")
        if contradiction_count > 0:
            lines.append(f"- intra_document_tensions: {contradiction_count}")
        if failed_bands > 0:
            lines.append(f"- failed_bands: {failed_bands}")
        if high_signal_cov is not None:
            lines.append(f"- high_signal_coverage: {float(high_signal_cov or 0.0):.2%}")
        if appendix_required:
            lines.append(f"- appendix_coverage: {'covered' if appendix_covered else 'missing'}")
        if short_summary:
            lines.append(f"- summary: {short_summary}")
        if themes:
            lines.append(f"- themes: {', '.join(str(x) for x in themes[:6])}")
        if tones:
            lines.append(f"- tone: {', '.join(str(x) for x in tones[:6])}")
        if framings:
            lines.append(f"- framing: {', '.join(str(x) for x in framings[:6])}")
        if opportunities:
            lines.append(f"- opportunities: {' | '.join(str(x) for x in opportunities[:4])}")
        if risks:
            lines.append(f"- risks: {' | '.join(str(x) for x in risks[:4])}")
        lines.append("")
    return "\n".join(lines).strip()


def _build_campaign_intel_context_block(snapshot: Optional[Dict[str, Any]]) -> str:
    if not snapshot:
        return ""
    dominant_narratives = list(snapshot.get("dominant_narratives") or [])[:10]
    opportunities = list(snapshot.get("strategic_opportunities") or [])[:10]
    risks = list(snapshot.get("strategic_risks") or [])[:10]
    contradictions = _safe_json_loads(snapshot.get("contradiction_tensions_json"), [])
    sentiment_distribution = _safe_json_loads(snapshot.get("sentiment_distribution_json"), {})
    tone_distribution = _safe_json_loads(snapshot.get("tone_distribution_json"), {})
    framing_distribution = _safe_json_loads(snapshot.get("framing_distribution_json"), {})
    lines: List[str] = ["=== CAMPAIGN-WIDE INTELLIGENCE ==="]
    lines.append(f"- snapshot_version: {snapshot.get('version')}")
    docs_total = int(snapshot.get("docs_total") or 0)
    docs_analyzed = int(snapshot.get("docs_analyzed") or 0)
    docs_failed = int(snapshot.get("docs_failed") or 0)
    coverage_ratio = float(snapshot.get("doc_intel_coverage_ratio") or 0.0)
    completeness_status = str(snapshot.get("input_completeness_status") or "").strip().lower()
    completeness_note = str(snapshot.get("input_completeness_note") or "").strip()
    quality_status = str(snapshot.get("input_quality_status") or "").strip().lower()
    quality_note = str(snapshot.get("input_quality_note") or "").strip()
    degraded_docs = int(snapshot.get("doc_intel_degraded_docs") or 0)
    full_fidelity_docs = int(snapshot.get("doc_intel_full_fidelity_docs") or 0)
    if docs_total > 0:
        lines.append(
            f"- input_coverage: analyzed {docs_analyzed}/{docs_total} indexed docs "
            f"(coverage={coverage_ratio:.2%}, missing={docs_failed})"
        )
    if completeness_status:
        lines.append(f"- input_completeness_status: {completeness_status}")
    if completeness_note:
        lines.append(f"- input_completeness_note: {completeness_note}")
    if docs_analyzed > 0:
        lines.append(
            f"- input_fidelity_mix: full_fidelity_docs={full_fidelity_docs}, degraded_docs={degraded_docs}"
        )
    if quality_status:
        lines.append(f"- input_quality_status: {quality_status}")
    if quality_note:
        lines.append(f"- input_quality_note: {quality_note}")
    if completeness_status in {"partial", "none"}:
        lines.append(
            "- caution: campaign intelligence is based on incomplete document-intelligence inputs; "
            "treat conclusions as provisional until coverage improves."
        )
    if quality_status in {"mixed_fidelity", "degraded_fidelity"}:
        lines.append(
            "- caution: some document-intelligence inputs were reduced-context; validate pivotal conclusions "
            "against raw source evidence before strategic commitments."
        )
    if dominant_narratives:
        lines.append(f"- dominant_narratives: {', '.join(str(x) for x in dominant_narratives)}")
    if opportunities:
        lines.append(f"- strategic_opportunities: {' | '.join(str(x) for x in opportunities[:6])}")
    if risks:
        lines.append(f"- strategic_risks: {' | '.join(str(x) for x in risks[:6])}")
    if isinstance(sentiment_distribution, dict) and sentiment_distribution:
        lines.append(f"- sentiment_distribution: {sentiment_distribution}")
    if isinstance(tone_distribution, dict) and tone_distribution:
        lines.append(f"- tone_distribution: {tone_distribution}")
    if isinstance(framing_distribution, dict) and framing_distribution:
        lines.append(f"- framing_distribution: {framing_distribution}")
    if isinstance(contradictions, list) and contradictions:
        summaries = [
            str(item.get("contradiction_summary") or "").strip()
            for item in contradictions[:6]
            if isinstance(item, dict) and str(item.get("contradiction_summary") or "").strip()
        ]
        if summaries:
            lines.append(f"- contradiction_tensions: {' | '.join(summaries)}")
    return "\n".join(lines).strip()


def _extract_openai_response_text(response_json: Dict[str, Any]) -> str:
    output_text = response_json.get("output_text")
    if isinstance(output_text, str) and output_text.strip():
        return output_text
    output = response_json.get("output")
    if not isinstance(output, list):
        return ""
    for item in output:
        if not isinstance(item, dict):
            continue
        content = item.get("content")
        if not isinstance(content, list):
            continue
        for segment in content:
            if not isinstance(segment, dict):
                continue
            text = segment.get("text")
            if isinstance(text, str) and text.strip():
                return text
    return ""


def _validate_and_sanitize_quotes(
    response_text: str,
    private_results: List[Dict[str, Any]],
    global_results: List[Dict[str, Any]],
) -> str:
    settings = get_settings()
    corpus_parts: List[str] = []
    for result in [*private_results, *global_results]:
        payload = result.get("payload", {}) or {}
        text = _get_text_for_rag(payload, settings)
        if text:
            corpus_parts.append(text)
    corpus = "\n".join(corpus_parts)
    if not corpus.strip():
        return response_text

    invalid_found = False

    def _replace_invalid(match: Any) -> str:
        nonlocal invalid_found
        opening = match.group(1)
        quoted = match.group(2)
        closing = match.group(3)
        if quoted in corpus:
            return f"{opening}{quoted}{closing}"
        invalid_found = True
        return quoted

    sanitized = re.sub(r'(["“])([^"\n”]{3,})(["”])', _replace_invalid, response_text)
    if invalid_found:
        sanitized = (
            f"{sanitized}\n\n"
            "Note: Quote not found in sources; re-run with more context."
        )
    return sanitized


def _build_report_prompt(
    *,
    query: str,
    context: str,
    has_context: bool,
    mode: str,
    report_type: str,
    user_display_name: str,
) -> str:
    persona_context = get_persona_context()
    normalized_report_type = _normalize_report_type(report_type)
    report_type_block = _report_type_instruction(normalized_report_type)
    mode_instruction = (
        "Deep report mode: produce a thorough memo with synthesis across documents, "
        "contradictions, risks, strategic opportunities, and explicit recommendations."
        if mode == "deep"
        else "Normal report mode: produce a focused strategic brief with concise evidence and recommendations."
    )
    if normalized_report_type == "summary":
        mode_instruction = (
            "Deep report mode: produce a thorough neutral synthesis across documents with strong evidence grounding."
            if mode == "deep"
            else "Normal report mode: produce a focused neutral synthesis with concise evidence grounding."
        )
    output_format_block = (
        """- Title
- Narrative Overview
- Key Themes Across Documents
- Contradictions and Tensions
- Potential Opponent-Exploitable Vulnerabilities
- Evidence References
- Clarifying Questions (if needed)"""
        if normalized_report_type == "summary"
        else """- Title
- Executive Summary
- Evidence from Sources
- Strategic Analysis
- Risks and Opportunities
- Recommendations and Next Moves
- Clarifying Questions (if needed)"""
    )
    return f"""You are Riley, a senior campaign strategist and decision partner at RALLY.

PERSONALITY AND LEADERSHIP STANDARD:
- Be brilliant, kind, firm, and practical.
- Be supportive and empathetic, but never sycophantic.
- Do not optimize for pleasing the user.
- Be direct when assumptions or strategy are weak, unsupported, risky, or misaligned.

REPORT TASK:
- Produce a high-quality strategy report for {user_display_name}.
- Synthesize campaign documents first, then use global sources as supporting context.
- Use campaign intelligence artifacts by default when available.
- Use document intelligence summaries to accelerate synthesis, then verify claims against source evidence.
- Separate clearly:
  1) what sources explicitly say,
  2) synthesized intelligence (cross-document patterns, narratives, sentiment, framing, tensions),
  3) your strategic recommendation and execution plan.
- Compare sources, identify contradictions and patterns, and surface strategic implications.
- Ask clarifying questions at the end when unresolved uncertainty materially affects recommendations.
- Ground factual claims with source references in format `[[Source: Filename.ext]]`.
- If evidence is incomplete, be explicit about limits instead of pretending certainty.

STRATEGIC PERSONA RECOGNITION:
Refer to these Strategic Archetypes: {persona_context}

OUTPUT FORMAT:
{output_format_block}

{report_type_block}

{mode_instruction}

Context Data:
{context if has_context else "[No specific data found in files yet]"}

User Request:
{query}
"""


def _build_summary_text(report_body: str) -> str:
    cleaned = (report_body or "").strip()
    if not cleaned:
        return "Report completed."
    first_block = cleaned.split("\n\n")[0].strip()
    if len(first_block) <= 320:
        return first_block
    return f"{first_block[:317].rstrip()}..."


def _format_size(size_bytes: int) -> str:
    if size_bytes < 1024:
        return f"{size_bytes} B"
    if size_bytes < 1024 * 1024:
        return f"{size_bytes / 1024:.1f} KB"
    return f"{size_bytes / (1024 * 1024):.1f} MB"


def _split_report_sections(report_body: str) -> Dict[str, str]:
    """Parse markdown-ish report output into named sections."""
    text = (report_body or "").strip()
    sections: Dict[str, str] = {
        "title": "",
        "executive_summary": "",
        "evidence_from_sources": "",
        "strategic_analysis": "",
        "recommendation": "",
        "clarifying_questions": "",
    }
    if not text:
        return sections

    # Strip markdown heading markers and normalize keys.
    pattern = re.compile(
        r"(?:^|\n)#{1,4}\s*(Title|Executive Summary|Evidence from Sources|Strategic Analysis|Recommendation|Clarifying Questions)\s*\n",
        flags=re.IGNORECASE,
    )
    matches = list(pattern.finditer(text))
    if not matches:
        sections["executive_summary"] = text
        return sections

    # Any preface before first known heading is treated as title/summary preface.
    prefix = text[: matches[0].start()].strip()
    if prefix:
        first_line = prefix.splitlines()[0].strip()
        sections["title"] = first_line[:180]
        if len(prefix.splitlines()) > 1:
            sections["executive_summary"] = "\n".join(prefix.splitlines()[1:]).strip()

    for idx, match in enumerate(matches):
        section_name = match.group(1).strip().lower().replace(" ", "_")
        start = match.end()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        section_text = text[start:end].strip()
        if section_name in sections:
            sections[section_name] = section_text

    return sections


def _strip_markdown_prefix_markers(text: str) -> str:
    value = str(text or "").strip()
    value = re.sub(r"^\s{0,3}#{1,6}\s*", "", value)
    value = re.sub(r"^\s*[-*+]\s+", "", value)
    value = re.sub(r"^\s*\d+[.)]\s+", "", value)
    return value.strip()


def _add_inline_markdown_runs(paragraph: Any, text: str) -> None:
    value = str(text or "")
    if not value:
        return
    # Normalize links/code to readable text first.
    value = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", value)
    value = re.sub(r"`([^`]+)`", r"\1", value)

    token_pattern = re.compile(
        r"(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|__[^_]+__|\*[^*\n]+\*|_[^_\n]+_)"
    )
    parts = token_pattern.split(value)
    for part in parts:
        if not part:
            continue
        bold = False
        italic = False
        content = part
        if part.startswith("***") and part.endswith("***") and len(part) > 6:
            content = part[3:-3]
            bold = True
            italic = True
        elif part.startswith("**") and part.endswith("**") and len(part) > 4:
            content = part[2:-2]
            bold = True
        elif part.startswith("__") and part.endswith("__") and len(part) > 4:
            content = part[2:-2]
            bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            content = part[1:-1]
            italic = True
        elif part.startswith("_") and part.endswith("_") and len(part) > 2:
            content = part[1:-1]
            italic = True
        run = paragraph.add_run(content)
        if bold:
            run.bold = True
        if italic:
            run.italic = True


def _add_markdownish_paragraph(document: Any, text: str, *, style: Optional[str] = None) -> None:
    paragraph = document.add_paragraph(style=style) if style else document.add_paragraph()
    _add_inline_markdown_runs(paragraph, _strip_markdown_prefix_markers(text))


def _render_markdownish_block(document: Any, text: str, *, default_heading_level: int = 2) -> None:
    raw = str(text or "").replace("\r\n", "\n").replace("\r", "\n")
    if not raw.strip():
        return

    lines = raw.split("\n")
    paragraph_buffer: List[str] = []

    def flush_paragraph() -> None:
        if not paragraph_buffer:
            return
        merged = " ".join(line.strip() for line in paragraph_buffer if line.strip()).strip()
        paragraph_buffer.clear()
        if merged:
            _add_markdownish_paragraph(document, merged)

    for line in lines:
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            continue

        heading_match = re.match(r"^\s{0,3}(#{1,6})\s+(.+)$", line)
        if heading_match:
            flush_paragraph()
            level = min(4, max(1, default_heading_level - 1 + len(heading_match.group(1))))
            heading_text = _strip_markdown_prefix_markers(heading_match.group(2))
            document.add_heading(heading_text, level=level)
            continue

        bullet_match = re.match(r"^\s*[-*+]\s+(.+)$", line)
        if bullet_match:
            flush_paragraph()
            _add_markdownish_paragraph(document, bullet_match.group(1), style="List Bullet")
            continue

        numbered_match = re.match(r"^\s*\d+[.)]\s+(.+)$", line)
        if numbered_match:
            flush_paragraph()
            _add_markdownish_paragraph(document, numbered_match.group(1), style="List Number")
            continue

        paragraph_buffer.append(line)

    flush_paragraph()


def _build_sources_appendix(
    private_results: List[Dict[str, Any]],
    global_results: List[Dict[str, Any]],
) -> List[str]:
    lines: List[str] = []
    seen: set[str] = set()
    for scope, items in (("Campaign", private_results), ("Global", global_results)):
        for result in items:
            payload = result.get("payload", {}) or {}
            filename = str(payload.get("filename") or "Unknown")
            location = str(payload.get("location_value") or payload.get("location") or "Unknown location")
            key = f"{scope}|{filename}|{location}"
            if key in seen:
                continue
            seen.add(key)
            lines.append(f"{scope}: {filename} — {location}")
    return lines


def _generate_docx_bytes(
    *,
    report_job_id: str,
    tenant_id: str,
    user_id: str,
    report_type: str,
    title: str,
    query_text: str,
    report_body: str,
    generated_at_iso: str,
    sources_appendix: List[str],
) -> bytes:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is not installed") from exc

    sections = _split_report_sections(report_body)
    effective_title = _strip_markdown_prefix_markers(
        sections.get("title") or title or "Riley Strategy Report"
    ) or "Riley Strategy Report"
    document = Document()
    document.add_heading(effective_title, level=0)
    document.add_paragraph(f"Generated: {generated_at_iso}")
    document.add_paragraph(f"Tenant/Campaign Context: {tenant_id}")
    document.add_paragraph(f"Report Job ID: {report_job_id}")
    document.add_paragraph(f"Report Type: {report_type}")
    document.add_paragraph(f"Created By User: {user_id}")

    document.add_heading("User Request", level=1)
    _render_markdownish_block(document, query_text.strip() or "[No request provided]", default_heading_level=2)

    document.add_heading("Executive Summary", level=1)
    _render_markdownish_block(document, sections.get("executive_summary") or report_body, default_heading_level=2)

    document.add_heading("Evidence from Sources", level=1)
    _render_markdownish_block(
        document,
        sections.get("evidence_from_sources") or "No explicit evidence section was generated.",
        default_heading_level=2,
    )

    document.add_heading("Strategic Analysis", level=1)
    _render_markdownish_block(
        document,
        sections.get("strategic_analysis") or "No explicit strategic analysis section was generated.",
        default_heading_level=2,
    )

    document.add_heading("Recommendation", level=1)
    _render_markdownish_block(
        document,
        sections.get("recommendation") or "No explicit recommendation section was generated.",
        default_heading_level=2,
    )

    clarifying_questions = (sections.get("clarifying_questions") or "").strip()
    if clarifying_questions:
        document.add_heading("Clarifying Questions", level=1)
        _render_markdownish_block(document, clarifying_questions, default_heading_level=2)

    if sources_appendix:
        document.add_heading("Sources / Appendix", level=1)
        for line in sources_appendix:
            document.add_paragraph(line, style="List Bullet")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


async def _persist_report_docx_artifact(
    *,
    report_job_id: str,
    tenant_id: str,
    user_id: str,
    report_type: str,
    title: str,
    query_text: str,
    report_body: str,
    sources_appendix: List[str],
    graph: Optional[GraphService] = None,
) -> Tuple[str, str]:
    settings = get_settings()
    generated_at = datetime.now().isoformat()
    safe_title = re.sub(r"[^a-zA-Z0-9_-]+", "_", (title or "riley_report").strip()).strip("_")
    if not safe_title:
        safe_title = "riley_report"
    artifact_filename = f"{safe_title}_{report_job_id[:8]}.docx"
    object_name = f"reports/{tenant_id}/{report_job_id}/{artifact_filename}"
    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    logger.info(
        "report_docx_generation_started report_job_id=%s tenant_id=%s report_type=%s",
        report_job_id,
        tenant_id,
        report_type,
    )
    docx_bytes = _generate_docx_bytes(
        report_job_id=report_job_id,
        tenant_id=tenant_id,
        user_id=user_id,
        report_type=report_type,
        title=title,
        query_text=query_text,
        report_body=report_body,
        generated_at_iso=generated_at,
        sources_appendix=sources_appendix,
    )
    logger.info(
        "report_docx_generation_completed report_job_id=%s tenant_id=%s bytes=%s",
        report_job_id,
        tenant_id,
        len(docx_bytes),
    )
    if graph and await _check_report_job_cancellation(
        report_job_id=report_job_id,
        graph=graph,
        checkpoint="before_upload",
    ):
        raise RuntimeError("Report cancelled before upload")
    logger.info(
        "report_upload_started report_job_id=%s tenant_id=%s object_name=%s",
        report_job_id,
        tenant_id,
        object_name,
    )
    output_url = await StorageService.upload_bytes(
        object_name=object_name,
        data=docx_bytes,
        content_type=mime_type,
    )
    logger.info(
        "report_upload_completed report_job_id=%s tenant_id=%s object_name=%s",
        report_job_id,
        tenant_id,
        object_name,
    )

    output_file_id = str(uuid.uuid4())
    payload: Dict[str, Any] = {
        "record_type": "file",
        "filename": artifact_filename,
        "file_type": "docx",
        "type": "docx",
        "mime_type": mime_type,
        "url": output_url,
        "is_global": tenant_id == "global",
        "tags": ["riley_report", report_type],
        "size": _format_size(len(docx_bytes)),
        "size_bytes": len(docx_bytes),
        "upload_date": generated_at,
        "uploaded_at": generated_at,
        "ai_enabled": True,
        "raw_content": report_body,
        "cleaned_content": report_body,
        "content": report_body,
        "content_preview": _build_summary_text(report_body),
        "ingestion_status": "indexed",
        "ingestion_error": None,
        "extracted_char_count": len(report_body or ""),
        "chunk_count": 0,
        "embedding_model": settings.EMBEDDING_MODEL,
        "embedding_tokens_estimate": 0,
        "embedding_cost_estimate_usd": 0.0,
        "chunk_profiles": {"micro": 0, "macro": 0},
        "bm25_enabled": False,
        "preview_url": None,
        "preview_type": None,
        "preview_status": "not_requested",
        "preview_error": None,
        "ocr_enabled": False,
        "ocr_status": None,
        "source": "riley_report",
        "report_job_id": report_job_id,
        "report_type": report_type,
        "created_by_user_id": user_id,
        "tenant_id": tenant_id,
    }
    if tenant_id != "global":
        payload["client_id"] = tenant_id

    placeholder_vector = [0.0] * int(settings.EMBEDDING_DIM)
    target_collection = (
        settings.QDRANT_COLLECTION_TIER_1 if tenant_id == "global"
        else settings.QDRANT_COLLECTION_TIER_2
    )
    await vector_service.client.upsert(
        collection_name=target_collection,
        points=[PointStruct(id=output_file_id, vector=placeholder_vector, payload=payload)],
    )
    logger.info(
        "report_artifact_persisted report_job_id=%s tenant_id=%s output_file_id=%s collection=%s",
        report_job_id,
        tenant_id,
        output_file_id,
        target_collection,
    )
    return output_file_id, output_url


def _build_worker_payload(report_job_id: str) -> Dict[str, str]:
    return {"report_job_id": report_job_id}


async def _enqueue_report_cloud_task(report_job_id: str) -> None:
    settings = get_settings()
    if not settings.RILEY_REPORTS_USE_CLOUD_TASKS:
        return
    if not settings.GCP_PROJECT_ID or not settings.RILEY_REPORT_WORKER_URL:
        raise RuntimeError(
            "Cloud Tasks report worker is enabled but GCP_PROJECT_ID/RILEY_REPORT_WORKER_URL is missing"
        )

    payload = _build_worker_payload(report_job_id)
    logger.info("report_queue_dispatch_started report_job_id=%s queue=%s", report_job_id, settings.RILEY_REPORTS_TASKS_QUEUE)

    def _create_task_sync() -> None:
        client = tasks_v2.CloudTasksClient()
        parent = client.queue_path(
            settings.GCP_PROJECT_ID,
            settings.RILEY_REPORTS_TASKS_LOCATION,
            settings.RILEY_REPORTS_TASKS_QUEUE,
        )
        task_name = f"{parent}/tasks/{report_job_id}"
        headers = {"Content-Type": "application/json"}
        if settings.RILEY_REPORT_WORKER_TOKEN:
            headers["X-Riley-Report-Worker-Token"] = settings.RILEY_REPORT_WORKER_TOKEN
        http_request: Dict[str, Any] = {
            "http_method": tasks_v2.HttpMethod.POST,
            "url": settings.RILEY_REPORT_WORKER_URL,
            "headers": headers,
            "body": json.dumps(payload).encode("utf-8"),
        }
        if settings.RILEY_REPORTS_TASKS_SERVICE_ACCOUNT_EMAIL:
            http_request["oidc_token"] = {
                "service_account_email": settings.RILEY_REPORTS_TASKS_SERVICE_ACCOUNT_EMAIL,
                "audience": settings.RILEY_REPORT_WORKER_URL,
            }
        task = {"name": task_name, "http_request": http_request}
        try:
            client.create_task(request={"parent": parent, "task": task})
        except AlreadyExists:
            return

    await run_in_threadpool(_create_task_sync)
    logger.info("report_queue_dispatch_completed report_job_id=%s queue=%s", report_job_id, settings.RILEY_REPORTS_TASKS_QUEUE)


async def create_report_job(
    *,
    graph: GraphService,
    tenant_id: str,
    user_id: str,
    query_text: str,
    mode: str,
    report_type: Optional[str] = None,
    title: Optional[str] = None,
    conversation_id: Optional[str] = None,
) -> Dict[str, Any]:
    settings = get_settings()
    report_job_id = str(uuid.uuid4())
    normalized_mode = _normalize_report_mode(mode)
    normalized_report_type = _normalize_report_type(report_type)
    job = await graph.create_riley_report_job(
        report_job_id=report_job_id,
        tenant_id=tenant_id,
        user_id=user_id,
        conversation_id=conversation_id,
        report_type=normalized_report_type,
        title=_derive_report_title(query_text, title),
        query_text=query_text,
        mode=normalized_mode,
    )
    logger.info(
        "report_job_created report_job_id=%s tenant_id=%s user_id=%s report_type=%s deep_mode=%s",
        report_job_id,
        tenant_id,
        user_id,
        normalized_report_type,
        normalized_mode == "deep",
    )
    try:
        if settings.RILEY_REPORTS_USE_CLOUD_TASKS:
            await _enqueue_report_cloud_task(report_job_id)
        else:
            # Non-durable fallback for local/dev environments.
            logger.info("report_inline_dispatch_started report_job_id=%s", report_job_id)
            asyncio.create_task(run_report_job(report_job_id=report_job_id, graph=graph))
            logger.info("report_inline_dispatch_completed report_job_id=%s", report_job_id)
    except Exception as exc:
        await graph.update_riley_report_job(
            report_job_id=report_job_id,
            tenant_id=tenant_id,
            user_id=user_id,
            status="failed",
            completed_at=datetime.now().isoformat(),
            error_message=f"Failed to enqueue report job: {type(exc).__name__}",
            failure_stage="queue_dispatch",
            failure_code=type(exc).__name__,
            failure_detail=str(exc)[:500],
        )
        logger.error(
            "report_job_dispatch_failed report_job_id=%s tenant_id=%s error_type=%s error_message=%s",
            report_job_id,
            tenant_id,
            type(exc).__name__,
            str(exc)[:280],
        )
        raise
    return job


async def _embed_query_text(content: str) -> List[float]:
    settings = get_settings()
    model_name = settings.EMBEDDING_MODEL

    def _embed_sync() -> List[float]:
        client = get_genai_client()
        response = client.models.embed_content(model=model_name, contents=content)
        if not response.embeddings:
            raise RuntimeError("Embedding response did not include vectors")
        values = response.embeddings[0].values
        if not isinstance(values, list):
            raise RuntimeError("Embedding response shape was invalid")
        return values

    try:
        return await run_in_threadpool(_embed_sync)
    except Exception as exc:
        raise RuntimeError(f"Query embedding failed: {exc}") from exc


async def _call_openai_report_model(*, prompt: str, model_name: str, timeout_seconds: int) -> str:
    settings = get_settings()
    if not settings.OPENAI_API_KEY:
        raise RuntimeError("OPENAI_API_KEY is not configured")
    try:
        import httpx  # type: ignore
    except ImportError as exc:
        raise RuntimeError("httpx dependency is not installed") from exc

    payload = {"model": model_name, "input": prompt}
    timeout = httpx.Timeout(timeout_seconds)
    async with httpx.AsyncClient(timeout=timeout) as client:
        response = await client.post(
            "https://api.openai.com/v1/responses",
            headers={
                "Authorization": f"Bearer {settings.OPENAI_API_KEY}",
                "Content-Type": "application/json",
            },
            json=payload,
        )
        response.raise_for_status()
        text = _extract_openai_response_text(response.json())
        if not text:
            raise RuntimeError("OpenAI report response did not contain text output")
        return text


def _is_retryable_report_error(exc: Exception) -> bool:
    name = type(exc).__name__.lower()
    message = str(exc).lower()
    if "timeout" in name:
        return True
    if "ratelimit" in name or "toomanyrequests" in name:
        return True
    return (
        "readtimeout" in message
        or "timed out" in message
        or "temporarily unavailable" in message
        or "status code: 429" in message
        or "status code: 500" in message
        or "status code: 502" in message
        or "status code: 503" in message
        or "status code: 504" in message
    )


def _sample_results_for_coverage(results: List[Dict[str, Any]], target: int) -> List[Dict[str, Any]]:
    if target <= 0 or not results:
        return []
    if target >= len(results):
        return list(results)
    if target == 1:
        return [results[0]]
    picks: List[Dict[str, Any]] = []
    seen: set[int] = set()
    for i in range(target):
        idx = round(i * (len(results) - 1) / (target - 1))
        if idx in seen:
            continue
        picks.append(results[idx])
        seen.add(idx)
    return picks


async def _check_report_job_cancellation(
    *,
    report_job_id: str,
    graph: GraphService,
    checkpoint: str,
) -> bool:
    job = await graph.get_riley_report_job_for_worker(report_job_id=report_job_id)
    if not job:
        logger.warning(
            "report_job_cancelled report_job_id=%s checkpoint=%s status=missing",
            report_job_id,
            checkpoint,
        )
        return True
    status = str(job.get("status") or "").lower()
    if status in {"cancelled", "deleted"}:
        logger.info(
            "report_job_cancelled report_job_id=%s checkpoint=%s status=%s",
            report_job_id,
            checkpoint,
            status,
        )
        return True
    return False


async def _retrieve_report_context(
    *,
    graph: Optional[GraphService],
    tenant_id: str,
    query_text: str,
) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]], str, List[str]]:
    settings = get_settings()
    query_vector = await _embed_query_text(query_text)

    rerank_candidates_limit = max(80, int(settings.RERANK_CANDIDATES), 100)
    rerank_top_k = max(20, int(settings.RERANK_TOP_K))
    private_limit = rerank_candidates_limit
    global_limit = rerank_candidates_limit
    private_results: List[Dict[str, Any]] = []
    global_results: List[Dict[str, Any]] = []
    graph_results = ""
    file_manifest: List[str] = []

    if tenant_id != "global":
        try:
            file_list = await vector_service.list_tenant_files(
                collection_name=settings.QDRANT_COLLECTION_TIER_2,
                tenant_id=tenant_id,
                limit=120,
            )
            file_manifest = [f.get("filename", "Unknown") for f in file_list if f.get("filename")]
        except Exception:
            file_manifest = []

    tasks: List[Any] = []
    if graph:
        tasks.append(graph.search_campaigns_fuzzy(query_text))
    if tenant_id != "global":
        private_filter = Filter(
            must=[
                FieldCondition(key="client_id", match=MatchValue(value=tenant_id)),
                FieldCondition(key="ai_enabled", match=MatchValue(value=True)),
            ],
            must_not=[FieldCondition(key="record_type", match=MatchValue(value="file"))],
        )
        if settings.HYBRID_SEARCH_ENABLED:
            tasks.append(
                vector_service.hybrid_search_research(
                    collection_name=settings.QDRANT_COLLECTION_TIER_2,
                    query_text=query_text,
                    query_embedding=query_vector,
                    tenant_filter=private_filter,
                    limit=private_limit,
                )
            )
        else:
            tasks.append(
                vector_service.search_silo(
                    collection_name=settings.QDRANT_COLLECTION_TIER_2,
                    query_vector=query_vector,
                    tenant_id=tenant_id,
                    limit=private_limit,
                    require_ai_enabled=True,
                )
            )

    global_filter = Filter(
        must=[FieldCondition(key="is_global", match=MatchValue(value=True))],
        must_not=[FieldCondition(key="record_type", match=MatchValue(value="file"))],
    )
    if settings.HYBRID_SEARCH_ENABLED:
        tasks.append(
            vector_service.hybrid_search_research(
                collection_name=settings.QDRANT_COLLECTION_TIER_1,
                query_text=query_text,
                query_embedding=query_vector,
                tenant_filter=global_filter,
                limit=global_limit,
            )
        )
    else:
        tasks.append(
            vector_service.search_global(
                collection_name=settings.QDRANT_COLLECTION_TIER_1,
                query_vector=query_vector,
                limit=global_limit,
                filter=global_filter,
            )
        )

    results = await asyncio.gather(*tasks, return_exceptions=True)
    idx = 0
    if graph:
        graph_result = results[idx]
        idx += 1
        if isinstance(graph_result, str):
            graph_results = graph_result
    if tenant_id != "global":
        private_result = results[idx]
        idx += 1
        if isinstance(private_result, list):
            private_results = private_result
    global_result = results[idx] if idx < len(results) else []
    if isinstance(global_result, list):
        global_results = global_result

    if settings.RERANK_ENABLED:
        combined = private_results + global_results
        ranked_ids = await rerank_candidates(
            query=query_text,
            candidates=combined,
            top_k=rerank_top_k,
        )
        if ranked_ids:
            private_results, global_results = _apply_rerank_order(
                private_results=private_results,
                global_results=global_results,
                ranked_ids=ranked_ids,
            )
            private_results = private_results[:rerank_top_k]
            remaining = max(0, rerank_top_k - len(private_results))
            global_results = global_results[:remaining]
        else:
            private_results = private_results[: min(private_limit, rerank_top_k)]
            global_results = global_results[: min(global_limit, rerank_top_k)]
    else:
        private_results = private_results[:private_limit]
        global_results = global_results[:global_limit]

    return private_results, global_results, graph_results, file_manifest


async def run_report_job(*, report_job_id: str, graph: GraphService) -> None:
    settings = get_settings()
    job = await graph.get_riley_report_job_for_worker(report_job_id=report_job_id)
    if not job:
        logger.warning("riley_report_worker_job_missing report_job_id=%s", report_job_id)
        return

    tenant_id = str(job.get("tenant_id") or "")
    user_id = str(job.get("user_id") or "")
    query_text = str(job.get("query") or "").strip()
    report_type = _normalize_report_type(str(job.get("report_type") or "strategy_memo"))
    title = str(job.get("title") or "Riley Strategy Report").strip() or "Riley Strategy Report"
    mode = _normalize_report_mode(str(job.get("mode") or "deep"))
    primary_model = settings.RILEY_GEMINI_MODEL
    fallback_model = settings.RILEY_OPENAI_FALLBACK_MODEL
    generation_model_used = primary_model

    logger.info(
        "report_job_worker_started report_job_id=%s tenant_id=%s report_type=%s model=%s deep_mode=%s",
        report_job_id,
        tenant_id,
        report_type,
        primary_model,
        mode == "deep",
    )

    if not query_text:
        await graph.update_riley_report_job(
            report_job_id=report_job_id,
            tenant_id=tenant_id,
            user_id=user_id,
            status="failed",
            completed_at=datetime.now().isoformat(),
            error_message="Report query text is empty",
            failure_stage="validation",
            failure_code="EMPTY_QUERY",
            failure_detail="Report query text is empty",
        )
        return

    started_at = datetime.now().isoformat()
    await graph.update_riley_report_job(
        report_job_id=report_job_id,
        tenant_id=tenant_id,
        user_id=user_id,
        status="processing",
        started_at=started_at,
        error_message=None,
        failure_stage=None,
        failure_code=None,
        failure_detail=None,
        generation_model=primary_model,
    )

    stage = "retrieval"
    report_body_for_debug: Optional[str] = None
    summary_for_debug: Optional[str] = None
    retrieval_doc_count = 0
    retrieval_chunk_count = 0
    used_context_chars = 0
    used_attempt_count = 0
    used_fidelity_level = "full"
    used_context_strategy = "full_context"
    context_reduction_applied = False

    try:
        if await _check_report_job_cancellation(
            report_job_id=report_job_id,
            graph=graph,
            checkpoint="before_retrieval",
        ):
            return
        logger.info(
            "report_retrieval_started report_job_id=%s tenant_id=%s report_type=%s deep_mode=%s",
            report_job_id,
            tenant_id,
            report_type,
            mode == "deep",
        )
        private_results, global_results, graph_results, file_manifest = await _retrieve_report_context(
            graph=graph,
            tenant_id=tenant_id,
            query_text=query_text,
        )
        retrieval_chunk_count = len(private_results) + len(global_results)
        unique_docs = {
            str((item.get("payload") or {}).get("parent_file_id") or item.get("id") or "")
            for item in [*private_results, *global_results]
            if str((item.get("payload") or {}).get("parent_file_id") or item.get("id") or "").strip()
        }
        retrieval_doc_count = len(unique_docs)
        logger.info(
            "report_retrieval_completed report_job_id=%s tenant_id=%s retrieval_doc_count=%s retrieval_chunk_count=%s graph_chars=%s",
            report_job_id,
            tenant_id,
            retrieval_doc_count,
            retrieval_chunk_count,
            len(graph_results or ""),
        )

        intelligence_blocks: List[str] = []
        try:
            private_doc_intel = await _load_parent_document_intelligence(
                collection_name=settings.QDRANT_COLLECTION_TIER_2,
                results=private_results,
            )
            global_doc_intel = await _load_parent_document_intelligence(
                collection_name=settings.QDRANT_COLLECTION_TIER_1,
                results=global_results,
            )
            campaign_snapshot = await graph.get_latest_riley_campaign_intelligence_snapshot(
                tenant_id=tenant_id
            )
            doc_intel_block = _build_doc_intel_context_block([*private_doc_intel, *global_doc_intel])
            if doc_intel_block:
                intelligence_blocks.append(doc_intel_block)
            campaign_block = _build_campaign_intel_context_block(campaign_snapshot)
            if campaign_block:
                intelligence_blocks.append(campaign_block)
        except Exception as intelligence_exc:
            logger.warning(
                "report_intelligence_context_unavailable report_job_id=%s tenant_id=%s error_type=%s error_message=%s",
                report_job_id,
                tenant_id,
                type(intelligence_exc).__name__,
                str(intelligence_exc)[:220],
            )

        retry_attempts = max(0, int(settings.RILEY_REPORT_RETRY_ATTEMPTS))
        max_attempts = retry_attempts + 1
        backoff_seconds = max(0.25, float(settings.RILEY_REPORT_RETRY_BACKOFF_SECONDS))
        base_timeout = max(45, int(settings.RILEY_REPORT_TIMEOUT_SECONDS))
        max_timeout = max(base_timeout, int(settings.RILEY_REPORT_MAX_TIMEOUT_SECONDS))
        base_context_chars = max(8000, int(settings.RILEY_REPORT_MAX_CONTEXT_CHARS))
        attempt_plan = [
            {"context_scale": 1.00, "timeout_scale": 1.00, "fidelity": "full", "strategy": "full_context"},
            {"context_scale": 0.85, "timeout_scale": 1.35, "fidelity": "moderately_reduced", "strategy": "coverage_preserving_85"},
            {"context_scale": 0.70, "timeout_scale": 1.80, "fidelity": "heavily_reduced", "strategy": "coverage_preserving_70"},
            {"context_scale": 0.55, "timeout_scale": 2.30, "fidelity": "heavily_reduced", "strategy": "coverage_preserving_55"},
            {"context_scale": 0.40, "timeout_scale": 2.90, "fidelity": "survival_mode", "strategy": "coverage_preserving_40"},
        ]

        report_body: Optional[str] = None
        last_exc: Optional[Exception] = None
        selected_private_results: List[Dict[str, Any]] = private_results
        selected_global_results: List[Dict[str, Any]] = global_results

        for attempt_idx in range(max_attempts):
            if await _check_report_job_cancellation(
                report_job_id=report_job_id,
                graph=graph,
                checkpoint="before_generation_attempt",
            ):
                return
            plan = attempt_plan[min(attempt_idx, len(attempt_plan) - 1)]
            context_scale = float(plan["context_scale"])
            context_max_chars = max(6000, int(round(base_context_chars * context_scale)))
            per_result_chars = max(900, int(round(2200 * context_scale)))
            p_target = max(6, int(round(len(private_results) * context_scale))) if private_results else 0
            g_target = max(4, int(round(len(global_results) * context_scale))) if global_results else 0
            selected_private_results = _sample_results_for_coverage(private_results, p_target)
            selected_global_results = _sample_results_for_coverage(global_results, g_target)

            context = _format_rag_context(
                selected_private_results,
                selected_global_results,
                graph_results,
                file_manifest,
                max_text_chars_per_result=per_result_chars,
            )
            if intelligence_blocks:
                intelligence_text = "\n\n".join(intelligence_blocks)
                if len(intelligence_text) > max(3000, int(context_max_chars * 0.4)):
                    intelligence_text = f"{intelligence_text[: max(3000, int(context_max_chars * 0.4))].rstrip()}…"
                context = f"{context}\n\n{intelligence_text}" if context else intelligence_text
            if len(context) > context_max_chars:
                context = f"{context[:context_max_chars].rstrip()}…"
                context_reduction_applied = True
            used_context_chars = len(context)
            has_context = bool(context.strip())
            prompt = _build_report_prompt(
                query=query_text,
                context=context,
                has_context=has_context,
                mode=mode,
                report_type=report_type,
                user_display_name="there",
            )
            timeout_seconds = min(
                max_timeout,
                max(base_timeout, int(round(base_timeout * float(plan["timeout_scale"])))),
            )
            used_fidelity_level = str(plan["fidelity"])
            used_context_strategy = str(plan["strategy"])
            used_attempt_count = attempt_idx + 1

            logger.info(
                "report_generation_started report_job_id=%s tenant_id=%s report_type=%s model=%s deep_mode=%s "
                "attempt=%s/%s timeout_s=%s retrieval_doc_count=%s context_chunks_included=%s context_chars=%s context_tokens_est=%s fidelity=%s",
                report_job_id,
                tenant_id,
                report_type,
                primary_model,
                mode == "deep",
                attempt_idx + 1,
                max_attempts,
                timeout_seconds,
                retrieval_doc_count,
                len(selected_private_results) + len(selected_global_results),
                used_context_chars,
                estimate_tokens(context),
                used_fidelity_level,
            )
            try:
                report_body = await generate_text_with_gemini(
                    prompt=prompt,
                    model_name=primary_model,
                    timeout_seconds=timeout_seconds,
                )
                generation_model_used = primary_model
                logger.info(
                    "report_generation_completed report_job_id=%s tenant_id=%s attempt=%s/%s model=%s",
                    report_job_id,
                    tenant_id,
                    attempt_idx + 1,
                    max_attempts,
                    generation_model_used,
                )
                break
            except Exception as exc:
                failure = classify_gemini_generation_failure(exc)
                logger.warning(
                    "gemini_generation_failed subsystem=%s tenant_id=%s report_job_id=%s primary_provider=%s fallback_provider=%s primary_model=%s fallback_model=%s error_type=%s http_status=%s provider_error_code=%s provider_error_type=%s fallback_eligible=%s",
                    "report",
                    tenant_id,
                    report_job_id,
                    "gemini",
                    "openai",
                    primary_model,
                    fallback_model,
                    failure.error_type,
                    failure.http_status,
                    failure.provider_error_code,
                    failure.provider_error_type,
                    failure.fallback_eligible,
                )
                if failure.fallback_eligible:
                    logger.warning(
                        "gemini_generation_fallback_to_openai subsystem=%s tenant_id=%s report_job_id=%s primary_provider=%s fallback_provider=%s primary_model=%s fallback_model=%s error_type=%s http_status=%s provider_error_code=%s provider_error_type=%s fallback_eligible=%s",
                        "report",
                        tenant_id,
                        report_job_id,
                        "gemini",
                        "openai",
                        primary_model,
                        fallback_model,
                        failure.error_type,
                        failure.http_status,
                        failure.provider_error_code,
                        failure.provider_error_type,
                        True,
                    )
                    try:
                        report_body = await _call_openai_report_model(
                            prompt=prompt,
                            model_name=fallback_model,
                            timeout_seconds=timeout_seconds,
                        )
                        generation_model_used = fallback_model
                        logger.info(
                            "provider_fallback_succeeded subsystem=%s tenant_id=%s report_job_id=%s primary_provider=%s fallback_provider=%s primary_model=%s fallback_model=%s",
                            "report",
                            tenant_id,
                            report_job_id,
                            "gemini",
                            "openai",
                            primary_model,
                            fallback_model,
                        )
                        logger.info(
                            "report_generation_completed report_job_id=%s tenant_id=%s attempt=%s/%s model=%s",
                            report_job_id,
                            tenant_id,
                            attempt_idx + 1,
                            max_attempts,
                            generation_model_used,
                        )
                        break
                    except Exception as fallback_exc:
                        fallback_failure = classify_openai_generation_failure(fallback_exc)
                        logger.warning(
                            "provider_fallback_failed subsystem=%s tenant_id=%s report_job_id=%s primary_provider=%s fallback_provider=%s primary_model=%s fallback_model=%s error_type=%s http_status=%s provider_error_code=%s provider_error_type=%s",
                            "report",
                            tenant_id,
                            report_job_id,
                            "gemini",
                            "openai",
                            primary_model,
                            fallback_model,
                            fallback_failure.error_type,
                            fallback_failure.http_status,
                            fallback_failure.provider_error_code,
                            fallback_failure.provider_error_type,
                        )
                        exc = fallback_exc
                last_exc = exc
                retryable = _is_retryable_report_error(exc) and attempt_idx < retry_attempts
                logger.warning(
                    "report_generation_attempt_failed report_job_id=%s tenant_id=%s attempt=%s/%s retryable=%s "
                    "error_type=%s error_message=%s",
                    report_job_id,
                    tenant_id,
                    attempt_idx + 1,
                    max_attempts,
                    retryable,
                    type(exc).__name__,
                    str(exc)[:260],
                )
                if not retryable:
                    raise
                await asyncio.sleep(backoff_seconds * (2 ** attempt_idx))

        if report_body is None:
            raise RuntimeError(
                f"Report generation retries exhausted: {type(last_exc).__name__ if last_exc else 'UnknownError'}"
            )

        stage = "post_generation"
        report_body = _validate_and_sanitize_quotes(report_body, selected_private_results, selected_global_results)
        summary_text = _build_summary_text(report_body)
        if used_fidelity_level != "full":
            summary_text = f"[Reduced coverage: {used_fidelity_level}] {summary_text}".strip()
        report_body_for_debug = report_body
        summary_for_debug = summary_text

        stage = "docx_generation"
        if await _check_report_job_cancellation(
            report_job_id=report_job_id,
            graph=graph,
            checkpoint="before_docx_generation",
        ):
            return
        sources_appendix = _build_sources_appendix(selected_private_results, selected_global_results)
        output_file_id, output_url = await _persist_report_docx_artifact(
            report_job_id=report_job_id,
            tenant_id=tenant_id,
            user_id=user_id,
            report_type=report_type,
            title=title,
            query_text=query_text,
            report_body=report_body,
            sources_appendix=sources_appendix,
            graph=graph,
        )

        stage = "persistence"
        if await _check_report_job_cancellation(
            report_job_id=report_job_id,
            graph=graph,
            checkpoint="before_final_persistence",
        ):
            return
        logger.info(
            "report_persistence_started report_job_id=%s tenant_id=%s output_file_id=%s",
            report_job_id,
            tenant_id,
            output_file_id,
        )
        await graph.update_riley_report_job(
            report_job_id=report_job_id,
            tenant_id=tenant_id,
            user_id=user_id,
            status="complete",
            completed_at=datetime.now().isoformat(),
            error_message=None,
            output_file_id=output_file_id,
            output_url=output_url,
            summary_text=summary_text,
            report_body=report_body,
            report_fidelity_level=used_fidelity_level,
            report_context_reduction_applied=context_reduction_applied,
            report_context_strategy=used_context_strategy,
            retrieval_doc_count=retrieval_doc_count,
            retrieval_chunk_count=len(selected_private_results) + len(selected_global_results),
            context_chars_included=used_context_chars,
            generation_model=generation_model_used,
            generation_attempts_used=used_attempt_count,
            failure_stage=None,
            failure_code=None,
            failure_detail=None,
        )
        logger.info(
            "report_persistence_completed report_job_id=%s tenant_id=%s final_status=complete fidelity=%s",
            report_job_id,
            tenant_id,
            used_fidelity_level,
        )
    except Exception as exc:
        error_type = type(exc).__name__
        error_message = str(exc)[:500]
        human_error = f"{stage}: {error_type}: {error_message}"
        logger.exception(
            "riley_report_job_failed report_job_id=%s tenant_id=%s report_type=%s model=%s deep_mode=%s "
            "stage=%s error_type=%s error_message=%s",
            report_job_id,
            tenant_id,
            report_type,
            generation_model_used,
            mode == "deep",
            stage,
            error_type,
            error_message,
        )
        await graph.update_riley_report_job(
            report_job_id=report_job_id,
            tenant_id=tenant_id,
            user_id=user_id,
            status="failed",
            completed_at=datetime.now().isoformat(),
            error_message=human_error[:1500],
            summary_text=summary_for_debug,
            report_body=report_body_for_debug,
            report_fidelity_level=used_fidelity_level,
            report_context_reduction_applied=context_reduction_applied,
            report_context_strategy=used_context_strategy,
            retrieval_doc_count=retrieval_doc_count or None,
            retrieval_chunk_count=retrieval_chunk_count or None,
            context_chars_included=used_context_chars or None,
            generation_model=generation_model_used,
            generation_attempts_used=used_attempt_count or None,
            failure_stage=stage,
            failure_code=error_type,
            failure_detail=error_message,
        )
